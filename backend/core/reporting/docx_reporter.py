"""
core/reporting/docx_reporter.py

Genera il report Word (.docx) da un record EmailAnalysis.
Struttura:
  1. Executive Summary
  2. Email Metadata
  3. Indicatori Tecnici (Header)
  4. Analisi Contenuto (Body + URL)
  5. Allegati
  6. Reputazione
  7. Valutazione del Rischio
  8. Note Manuali (sezione editabile)
"""

from pathlib import Path
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils.config import settings


RISK_COLORS = {
    "low": RGBColor(0x2E, 0x7D, 0x32),       # verde
    "medium": RGBColor(0xF5, 0x7F, 0x17),     # arancione
    "high": RGBColor(0xD3, 0x2F, 0x2F),       # rosso
    "critical": RGBColor(0x6A, 0x1B, 0x9A),   # viola
}

SEVERITY_COLORS = {
    "info": RGBColor(0x01, 0x57, 0x9B),
    "low": RGBColor(0x33, 0x69, 0x1E),
    "medium": RGBColor(0xF5, 0x7F, 0x17),
    "high": RGBColor(0xD3, 0x2F, 0x2F),
    "critical": RGBColor(0x6A, 0x1B, 0x9A),
}


def _add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def _add_kv(doc: Document, key: str, value: str):
    """Aggiunge una riga chiave: valore."""
    p = doc.add_paragraph()
    run_key = p.add_run(f"{key}: ")
    run_key.bold = True
    run_key.font.size = Pt(10)
    run_val = p.add_run(str(value) if value else "N/A")
    run_val.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)


def _add_finding_row(doc: Document, severity: str, description: str, evidence: str = ""):
    p = doc.add_paragraph(style="List Bullet")
    color = SEVERITY_COLORS.get(severity.lower(), RGBColor(0, 0, 0))
    badge = p.add_run(f"[{severity.upper()}] ")
    badge.bold = True
    badge.font.color.rgb = color
    badge.font.size = Pt(9)
    desc_run = p.add_run(description)
    desc_run.font.size = Pt(9)
    if evidence:
        ev_run = p.add_run(f"\n  → {evidence[:300]}")
        ev_run.font.size = Pt(8)
        ev_run.font.color.rgb = RGBColor(0x78, 0x78, 0x78)


def generate_report(record, output_path: Path):
    doc = Document()

    # Stile base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ------------------------------------------------------------------ #
    # COPERTINA / TITOLO
    # ------------------------------------------------------------------ #
    title = doc.add_heading(f"EMLyzer v{settings.VERSION} – Report di Analisi", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Generato il: {datetime.now(timezone.utc).strftime('%d/%m/%Y %H:%M UTC')}").italic = True

    doc.add_paragraph()

    # ------------------------------------------------------------------ #
    # 1. EXECUTIVE SUMMARY
    # ------------------------------------------------------------------ #
    _add_heading(doc, "1. Executive Summary")

    risk_label = (record.risk_label or "unknown").lower()
    risk_score = record.risk_score or 0.0
    risk_color = RISK_COLORS.get(risk_label, RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.add_run("Rischio complessivo: ").bold = True
    r = p.add_run(f"{risk_label.upper()} ({risk_score:.1f}/100)")
    r.bold = True
    r.font.color.rgb = risk_color
    r.font.size = Pt(12)

    explanation = []
    if record.risk_explanation and "explanation" in record.risk_explanation:
        explanation = record.risk_explanation["explanation"]

    if explanation:
        doc.add_paragraph("Principali indicatori di rischio:", style="Normal").runs[0].bold = True
        for item in explanation[:8]:
            doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ------------------------------------------------------------------ #
    # 2. EMAIL METADATA
    # ------------------------------------------------------------------ #
    _add_heading(doc, "2. Email Metadata")
    _add_kv(doc, "Oggetto", record.mail_subject)
    _add_kv(doc, "Da (From)", record.mail_from)
    _add_kv(doc, "A (To)", record.mail_to)
    _add_kv(doc, "Data", record.mail_date)
    _add_kv(doc, "Message-ID", record.message_id)
    _add_kv(doc, "Return-Path", record.return_path)
    _add_kv(doc, "Reply-To", record.reply_to)
    _add_kv(doc, "X-Mailer", record.x_mailer)
    _add_kv(doc, "X-Originating-IP", record.x_originating_ip)
    _add_kv(doc, "X-Campaign-ID", record.x_campaign_id)
    _add_kv(doc, "SHA256 file", record.file_hash_sha256)

    doc.add_paragraph()
    _add_heading(doc, "Autenticazione", level=2)
    _add_kv(doc, "SPF", record.spf_result or "N/A")
    _add_kv(doc, "DKIM", record.dkim_result or "N/A")
    _add_kv(doc, "DMARC", record.dmarc_result or "N/A")

    doc.add_page_break()

    # ------------------------------------------------------------------ #
    # 3. INDICATORI TECNICI (HEADER)
    # ------------------------------------------------------------------ #
    _add_heading(doc, "3. Indicatori Tecnici – Header")

    hi = record.header_indicators or {}
    findings = hi.get("findings", [])
    if findings:
        for f in findings:
            _add_finding_row(doc, f.get("severity", "info"), f.get("description", ""), f.get("evidence", ""))
    else:
        doc.add_paragraph("Nessun indicatore header rilevato.")

    # Catena Received
    hops = hi.get("received_hops", [])
    if hops:
        doc.add_paragraph()
        _add_heading(doc, "Percorso SMTP (Received chain)", level=2)
        for hop in hops:
            doc.add_paragraph(
                f"Hop {hop.get('hop', '?')}: IP={hop.get('ip', 'N/A')} | "
                f"by={hop.get('by', 'N/A')} | {hop.get('timestamp', '')}",
                style="List Bullet",
            )

    doc.add_page_break()

    # ------------------------------------------------------------------ #
    # 4. ANALISI CONTENUTO (BODY + URL)
    # ------------------------------------------------------------------ #
    _add_heading(doc, "4. Analisi del Contenuto")

    bi = record.body_indicators or {}
    _add_kv(doc, "Pattern urgenza", str(bi.get("urgency_count", 0)))
    _add_kv(doc, "Call-to-action sospette", str(bi.get("phishing_cta_count", 0)))
    _add_kv(doc, "Keyword credenziali", str(bi.get("credential_keyword_count", 0)))
    _add_kv(doc, "Form HTML embedded", str(bi.get("forms_found", 0)))
    _add_kv(doc, "JavaScript", "Sì" if bi.get("js_found") else "No")
    _add_kv(doc, "Elementi HTML nascosti", str(bi.get("invisible_elements", 0)))

    body_findings = bi.get("findings", [])
    if body_findings:
        doc.add_paragraph()
        _add_heading(doc, "Finding Body", level=2)
        for f in body_findings:
            _add_finding_row(doc, f.get("category", "info"), f.get("description", ""), f.get("evidence", ""))

    obfuscated = bi.get("obfuscated_links", [])
    if obfuscated:
        doc.add_paragraph()
        _add_heading(doc, "Link Offuscati", level=2)
        for lnk in obfuscated[:20]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("Testo visibile: ").bold = True
            p.add_run(str(lnk.get("visible_text", ""))[:150])
            p.add_run(" → href reale: ").bold = True
            p.add_run(str(lnk.get("actual_href", ""))[:200])

    # URL Analysis
    ui = record.url_indicators or {}
    url_list = ui.get("urls", [])
    if url_list:
        doc.add_paragraph()
        _add_heading(doc, "Analisi URL", level=2)
        _add_kv(doc, "URL totali", str(ui.get("total_urls", 0)))
        _add_kv(doc, "URL ad alto rischio", str(ui.get("high_risk_count", 0)))
        for u in url_list[:30]:
            flags = []
            if u.get("is_ip_address") or u.get("is_ip"):
                flags.append("IP diretto")
            if u.get("is_shortener"):
                flags.append("URL shortener")
            if u.get("is_punycode"):
                flags.append("Punycode/IDN")
            flag_str = " | ".join(flags) if flags else "OK"
            doc.add_paragraph(
                f"[score: {u.get('risk_score', 0):.0f}] {(u.get('original_url') or u.get('url', ''))[:200]} ({flag_str})",
                style="List Bullet",
            )

    doc.add_page_break()

    # ------------------------------------------------------------------ #
    # 5. ALLEGATI
    # ------------------------------------------------------------------ #
    _add_heading(doc, "5. Allegati")

    ai = record.attachment_indicators or {}
    _add_kv(doc, "Allegati totali", str(ai.get("total_attachments", 0)))
    _add_kv(doc, "Allegati critici", str(ai.get("critical_count", 0)))

    for att in ai.get("attachments", []):
        doc.add_paragraph()
        _add_heading(doc, f"Allegato: {att.get('filename', 'N/A')}", level=2)
        _add_kv(doc, "Dimensione", f"{att.get('size_bytes', 0):,} bytes")
        _add_kv(doc, "MIME dichiarato", att.get("declared_mime", ""))
        _add_kv(doc, "MIME reale", att.get("real_mime", ""))
        _add_kv(doc, "SHA256", att.get("hash_sha256", ""))
        for f in att.get("findings", []):
            _add_finding_row(doc, f.get("severity", "info"), f.get("description", ""), f.get("evidence", ""))

    doc.add_page_break()

    # ------------------------------------------------------------------ #
    # 6. REPUTAZIONE
    # ------------------------------------------------------------------ #
    _add_heading(doc, "6. Reputazione")

    rep = record.reputation_results
    if rep:
        _add_kv(doc, "Score reputazione", f"{rep.get('reputation_score', 0):.1f}/100")
        _add_kv(doc, "Indicatori malevoli", str(rep.get("malicious_count", 0)))

        # Entità analizzate (disponibili da v0.3.4+)
        ea = rep.get("entities_analyzed")
        if ea:
            _add_kv(doc, "IP analizzati", str(ea.get("ips", 0)))
            _add_kv(doc, "URL analizzati", str(ea.get("urls", 0)))
            if ea.get("hashes", 0) > 0:
                _add_kv(doc, "Hash analizzati", str(ea.get("hashes", 0)))

        # Risultati per categoria — separati per leggibilità
        ip_results  = [r for r in rep.get("ip_results",   []) if not r.get("skipped")]
        url_results = [r for r in rep.get("url_results",  []) if not r.get("skipped")]
        hsh_results = [r for r in rep.get("hash_results", []) if not r.get("skipped")]

        # Servizi di sicurezza: mostra tutti i risultati malevoli + errori
        for section_label, results in [
            ("Risultati IP", ip_results),
            ("Risultati URL", url_results),
            ("Risultati Hash allegati", hsh_results),
        ]:
            if not results:
                continue
            doc.add_paragraph()
            _add_heading(doc, section_label, level=2)
            for r in results:
                # Servizi informativi (ASN, crt.sh, Redirect Chain) → mostra solo il dettaglio
                info_services = {"ASN Lookup", "crt.sh", "Redirect Chain", "Shodan InternetDB"}
                if r.get("source") in info_services:
                    if r.get("detail") and not r.get("error"):
                        doc.add_paragraph(
                            f"[{r.get('source')}] {r.get('entity', '')[:80]}: {r.get('detail', '')[:200]}",
                            style="List Bullet",
                        )
                    elif r.get("error"):
                        doc.add_paragraph(
                            f"[{r.get('source')}] {r.get('entity', '')[:80]}: errore — {r.get('error', '')[:100]}",
                            style="List Bullet",
                        )
                    continue

                status = "⚠ MALEVOLO" if r.get("is_malicious") else "pulito"
                detail = r.get("detail", "")[:200]
                err    = r.get("error", "")
                line   = f"[{r.get('source')}] {r.get('entity', '')[:100]} → {status}"
                if detail:
                    line += f" | {detail}"
                if err:
                    line += f" | ERRORE: {err[:80]}"
                p = doc.add_paragraph(line, style="List Bullet")
                if r.get("is_malicious"):
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)
    else:
        doc.add_paragraph(
            "Analisi di reputazione non eseguita. "
            "Clicca 'Avvia controllo reputazione' nella scheda Reputazione "
            "oppure usa POST /api/reputation/{job_id}."
        )

    doc.add_page_break()

    # ------------------------------------------------------------------ #
    # 7. VALUTAZIONE DEL RISCHIO
    # ------------------------------------------------------------------ #
    _add_heading(doc, "7. Valutazione del Rischio")

    p = doc.add_paragraph()
    p.add_run(f"Score finale: ").bold = True
    r = p.add_run(f"{risk_score:.1f}/100 – {risk_label.upper()}")
    r.font.color.rgb = risk_color
    r.bold = True

    if record.risk_explanation:
        contribs = record.risk_explanation.get("contributions", [])
        if isinstance(contribs, list):
            _add_heading(doc, "Contributo per modulo", level=2)
            for c in contribs:
                if isinstance(c, dict):
                    doc.add_paragraph(
                        f"{c.get('module', '').capitalize()}: raw={c.get('raw_score', 0):.1f} → "
                        f"pesato={c.get('weighted_score', 0):.1f}",
                        style="List Bullet",
                    )

    doc.add_page_break()

    # ------------------------------------------------------------------ #
    # 8. NOTE MANUALI (sezione editabile)
    # ------------------------------------------------------------------ #
    _add_heading(doc, "8. Note dell'Analista")

    if record.analyst_notes:
        doc.add_paragraph(record.analyst_notes)
    else:
        p = doc.add_paragraph("[Sezione editabile – inserire qui le osservazioni manuali dell'analista]")
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # Salva
    doc.save(str(output_path))